#!/usr/bin/env python3
"""Genera Manual_de_Usuario_MP.docx con la documentación actualizada."""

import os
from docx import Document
from docx.shared import Pt

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUTPUT = os.path.join(ROOT, "Manual_de_Usuario_MP.docx")


def add_bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


def create_manual():
    doc = Document()
    doc.add_heading("Manual de Usuario — Sistema MP (Pasina / SARC)", 0)

    # 1. Intro
    doc.add_heading("1. Introducción", level=1)
    doc.add_paragraph(
        "El Sistema MP permite el alta rápida de clientes, su validación crediticia "
        "(Nosis) e impositiva (ARCA), la exportación hacia Presea ERP y el análisis "
        "de gestión comercial mediante tableros de ventas."
    )

    # 2. Roles
    doc.add_heading("2. Roles y Permisos", level=1)
    add_bullet(doc, "Vendedor (permiso_alta): carga clientes nuevos desde la app (código ≥ 40.000).")
    add_bullet(doc, "Validador (permiso_validacion): aprueba, rechaza o valida clientes.")
    add_bullet(doc, "Análisis (permisos_av): accede a tableros de ventas y carga histórica.")

    # 3. Alta desde app
    doc.add_heading("3. Alta de Clientes desde la Aplicación", level=1)
    doc.add_paragraph(
        "Flujo: ingresar CUIT → consultar ARCA → completar datos comerciales → "
        "guardar en cola de pendientes. Los clientes reciben código ≥ 40.000 al exportarse."
    )
    add_bullet(doc, "Consulta ARCA: obtiene razón social, domicilio fiscal, actividad y tipo responsable.")
    add_bullet(doc, "Consulta por voz: dictado de datos complementarios (opcional).")
    add_bullet(doc, "Estados: Pendiente → Modificado → A Exportar → Exportado.")

    # 4. Validación
    doc.add_heading("4. Validación de Clientes", level=1)
    doc.add_heading("4.1. Clientes Pendientes", level=2)
    doc.add_paragraph(
        "Lista clientes cargados por vendedores (origen app, código ≥ 40.000). "
        "Al seleccionar uno se ejecuta análisis Nosis automático. Acciones: "
        "Guardar cambios, Marcar para Exportar, Rechazar."
    )
    doc.add_heading("4.2. Clientes Rechazados", level=2)
    doc.add_paragraph("Clientes cuyo alta fue rechazada. Puede recuperarlos a pendientes.")
    doc.add_heading("4.3. Clientes Altas desde Presea (NUEVO)", level=2)
    doc.add_paragraph(
        "Clientes dados de alta directamente en Presea ERP (código < 40.000). "
        "Se importan automáticamente desde CLIENTESPA.DBI al sincronizar el servidor Windows."
    )
    add_bullet(doc, "Buscar por código, nombre o CUIT en la grilla.")
    add_bullet(doc, "Validar ARCA: consulta/modifica datos impositivos faltantes.")
    add_bullet(doc, "Validar NOSIS: análisis crediticio completo con reporte PDF.")
    add_bullet(doc, "Opción de enviar cambios de vuelta a Presea (genera DBI + FTP).")

    # 5. Presea sync
    doc.add_heading("5. Sincronización con Presea (windows_sync.exe)", level=1)
    doc.add_paragraph("El sincronizador en Windows Server ejecuta cada ~15 minutos:")
    add_bullet(doc, "EXPORTA → FTP + Supabase: CLIENTESPA.DBI, CODIGOSMP.DBI, ramo.csv, ventas.dbi")
    add_bullet(doc, "FTP → IMPORTA: Clientes_web.dbi exportados desde la app")
    add_bullet(doc, "Clientes Presea (codigo < 40000) → tabla clientes_pendientes, origen=presea")

    # 6. Ventas
    doc.add_heading("6. Análisis de Gestión (Ventas)", level=1)
    doc.add_paragraph(
        "Los tableros leen datos de Supabase (tabla ventas). El archivo ventas.dbi "
        "se importa desde Presea via windows_sync.exe. Streamlit Cloud NO lee el DBI directamente."
    )
    add_bullet(doc, "Carga histórica manual: subir ventas.dbi en el expander del módulo (permiso AV).")

    # 7. DBI structure
    doc.add_heading("7. Estructura CLIENTESPA.DBI", level=1)
    doc.add_paragraph(
        "Ver documento técnico docs/CLIENTESPA_DBI_ESTRUCTURA.md. Campos principales: "
        "CODIGO, NOMBRE, CUIT, DOMICILIO, LOCALIDAD, PROVINCIA, RUBRO, TIPO_RESP, VENDEDOR, MEMO."
    )
    doc.add_paragraph("Regla: CODIGO < 40000 = Presea | CODIGO ≥ 40000 = aplicación web.")

    # 8. Deploy
    doc.add_heading("8. Puesta en Marcha", level=1)
    add_bullet(doc, "GitHub: grandijuanluis-max/Validacion_Alta_CLientes_MP")
    add_bullet(doc, "Streamlit Cloud: app.py + secrets (.streamlit/secrets.toml.example)")
    add_bullet(doc, "Supabase: ejecutar supabase_schema.sql + migraciones")
    add_bullet(doc, "Windows Server: compilar windows_sync.exe + Task Scheduler")

    # 9. Troubleshooting
    doc.add_heading("9. Problemas Frecuentes", level=1)
    add_bullet(doc, "Sin clientes Presea en solapa: verificar windows_sync y migración SQL.")
    add_bullet(doc, "ARCA falla en Cloud: revisar certificados AFIP_CRT/AFIP_KEY en secrets.")
    add_bullet(doc, "Tableros vacíos: confirmar ventas.dbi importado en Supabase.")

    doc.save(OUTPUT)
    print(f"Manual generado: {OUTPUT}")


if __name__ == "__main__":
    create_manual()
